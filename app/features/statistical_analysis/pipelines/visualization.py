import os
import io
import base64
import matplotlib
# Force matplotlib to run in headless mode
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import seaborn as sns
import pandas as pd
import numpy as np
from typing import Dict, Any, List

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

from app.config import EXPORT_DIR

# Establish a consistent aesthetic theme
sns.set_theme(style="whitegrid")
plt.rcParams['font.family'] = 'sans-serif'
plt.rcParams['font.sans-serif'] = ['DejaVu Sans', 'Arial', 'Calibri']

# Premium color palette
primary_color = "#1E3A8A"  # Deep Blue
secondary_color = "#0EA5E9"  # Sky Blue
accent_color = "#F43F5E"  # Rose Red

def generate_plot(
    df: pd.DataFrame,
    model_type: str,
    independent_vars: List[str],
    dependent_var: str,
    group_var: str = None,
    time_col: str = None,
    time_series_results: Dict[str, Any] = None,
    analysis_results: Dict[str, Any] = None
) -> io.BytesIO:
    """
    Generates a statistical plot based on the model type.
    """
    fig, ax = plt.subplots(figsize=(8, 5.5), dpi=300)
    
    # 1. Time Series Plotting
    if time_series_results:
        ts = time_series_results
        # Sort by time_col if available
        if time_col and time_col in df.columns:
            df_ts = df.sort_values(by=time_col).copy()
        else:
            df_ts = df.copy()
            
        y_vals = pd.to_numeric(df_ts[dependent_var], errors='coerce').dropna().values
        x_vals = np.arange(len(y_vals))
        
        ax.plot(x_vals, y_vals, color=primary_color, linewidth=2, label="Data Historis")
        
        # Append forecasts
        forecasts = ts.get("forecasts_5_steps", [])
        if forecasts:
            x_fc = np.arange(len(y_vals), len(y_vals) + len(forecasts))
            # Connect the last point of history to the first point of forecast
            x_fc_conn = np.insert(x_fc, 0, len(y_vals) - 1)
            y_fc_conn = np.insert(np.array(forecasts), 0, y_vals[-1])
            
            ax.plot(x_fc_conn, y_fc_conn, color=accent_color, linestyle="--", linewidth=2, label="Peramalan (5 Langkah)")
            ax.scatter(x_fc, forecasts, color=accent_color, s=40, zorder=5)
            
        ax.set_title(f"Tren dan Hasil Peramalan: {dependent_var}", fontsize=14, pad=15, fontweight='bold', color='#1F2937')
        ax.set_xlabel("Waktu (Periode / Index)", fontsize=11, fontweight='semibold', labelpad=10)
        ax.set_ylabel(dependent_var, fontsize=11, fontweight='semibold', labelpad=10)
        ax.legend()
        
    # 2. Mediation Path Diagram
    elif model_type == "Path Analysis (Simple Mediation)":
        # Draw path diagram
        ax.axis('off')
        ax.set_xlim(0, 10)
        ax.set_ylim(0, 8)
        
        paths = analysis_results.get("paths", {})
        x_lbl = analysis_results.get("independent_var", "X")
        m_lbl = analysis_results.get("mediator_var", "M")
        y_lbl = analysis_results.get("dependent_var", "Y")
        
        a_coef = paths.get("path_a", {}).get("coefficient", 0.0)
        b_coef = paths.get("path_b", {}).get("coefficient", 0.0)
        c_prime = paths.get("path_c_prime", {}).get("coefficient", 0.0)
        
        # Draw boxes
        ax.text(2, 2, f"X\n({x_lbl})", ha='center', va='center', size=12, weight='bold',
                bbox=dict(boxstyle="round,pad=0.5", fc="#DBEAFE", ec="#1E3A8A", lw=1.5))
        ax.text(5, 6, f"M\n({m_lbl})", ha='center', va='center', size=12, weight='bold',
                bbox=dict(boxstyle="round,pad=0.5", fc="#D1FAE5", ec="#10B981", lw=1.5))
        ax.text(8, 2, f"Y\n({y_lbl})", ha='center', va='center', size=12, weight='bold',
                bbox=dict(boxstyle="round,pad=0.5", fc="#FEE2E2", ec="#EF4444", lw=1.5))
                
        # Draw arrows
        # Path a (X -> M)
        ax.annotate("", xy=(4.3, 5.5), xytext=(2.5, 2.5),
                    arrowprops=dict(arrowstyle="-|>", color="#1E293B", lw=2, mutation_scale=15))
        ax.text(3.1, 4.2, f"a = {a_coef:.3f}", ha='right', size=10, weight='semibold')
        
        # Path b (M -> Y)
        ax.annotate("", xy=(7.5, 2.5), xytext=(5.7, 5.5),
                    arrowprops=dict(arrowstyle="-|>", color="#1E293B", lw=2, mutation_scale=15))
        ax.text(6.9, 4.2, f"b = {b_coef:.3f}", ha='left', size=10, weight='semibold')
        
        # Path c' (X -> Y)
        ax.annotate("", xy=(7.1, 2.0), xytext=(2.9, 2.0),
                    arrowprops=dict(arrowstyle="-|>", color="#1E293B", lw=2, mutation_scale=15))
        ax.text(5.0, 1.6, f"c' = {c_prime:.3f}", ha='center', size=10, weight='semibold')
        
        ax.set_title("Diagram Jalur Analisis Mediasi (X \u2192 M \u2192 Y)", fontsize=14, pad=15, fontweight='bold', color='#1F2937')

    # 3. Heatmap for Multi-Variable Correlation
    elif "Correlation" in model_type and len(independent_vars) > 1:
        all_cols = [dependent_var] + independent_vars
        df_corr = df[all_cols].apply(pd.to_numeric, errors='coerce').corr()
        
        sns.heatmap(
            df_corr, 
            annot=True, 
            cmap="RdBu_r", 
            vmin=-1, 
            vmax=1, 
            square=True, 
            fmt=".3f", 
            linewidths=0.5, 
            cbar_kws={"shrink": 0.8}, 
            ax=ax
        )
        ax.set_title("Matriks Korelasi Koefisien", fontsize=14, pad=15, fontweight='bold', color='#1F2937')
        
    # 4. Boxplots/Swarmplots for Comparative Path
    elif "T-Test" in model_type or "Whitney" in model_type or "Wilcoxon" in model_type or "ANOVA" in model_type or "Kruskal" in model_type:
        sns.boxplot(
            data=df, 
            x=group_var, 
            y=dependent_var, 
            palette="Blues",
            width=0.5, 
            linewidth=1.5,
            fliersize=5, 
            ax=ax
        )
        if len(df) < 150:
            sns.stripplot(
                data=df, 
                x=group_var, 
                y=dependent_var, 
                color="#111827", 
                size=5, 
                jitter=0.15, 
                alpha=0.4, 
                ax=ax
            )
        ax.set_title(f"Perbandingan {dependent_var} Berdasarkan {group_var}", fontsize=14, pad=15, fontweight='bold', color='#1F2937')
        ax.set_xlabel(group_var, fontsize=11, fontweight='semibold', labelpad=10)
        ax.set_ylabel(dependent_var, fontsize=11, fontweight='semibold', labelpad=10)
        
    # 5. Scatter Plot with Trendline for Regression/Single Correlation
    else:
        primary_x = independent_vars[0] if independent_vars else df.columns[0]
        sns.scatterplot(
            data=df, 
            x=primary_x, 
            y=dependent_var, 
            color=primary_color, 
            alpha=0.7, 
            s=80, 
            edgecolor='w', 
            linewidth=0.5, 
            ax=ax
        )
        if len(df) > 1:
            try:
                sns.regplot(
                    data=df, 
                    x=primary_x, 
                    y=dependent_var, 
                    scatter=False, 
                    color=accent_color, 
                    line_kws={"linewidth": 2, "linestyle": "--"}, 
                    ax=ax
                )
            except Exception:
                pass
        ax.set_title(f"Hubungan antara {primary_x} dan {dependent_var}", fontsize=14, pad=15, fontweight='bold', color='#1F2937')
        ax.set_xlabel(primary_x, fontsize=11, fontweight='semibold', labelpad=10)
        ax.set_ylabel(dependent_var, fontsize=11, fontweight='semibold', labelpad=10)
        
    plt.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=300)
    buf.seek(0)
    plt.close(fig)
    return buf

def encode_image_base64(buf: io.BytesIO) -> str:
    return base64.b64encode(buf.getvalue()).decode('utf-8')

# Helper to clean up research titles for file naming
def clean_filename(title: str) -> str:
    cleaned = "".join([c if c.isalnum() or c in ("-", "_") else "_" for c in title])
    return cleaned[:50]

# --- WORD DOCX GENERATOR ---
def create_docx_report(
    file_id: str,
    research_title: str,
    filename: str,
    preprocessing_log: Dict[str, Any],
    normality_results: Dict[str, Any],
    analysis_results: Dict[str, Any],
    independent_vars: List[str],
    dependent_var: str,
    narrative: str,
    plot_buf: io.BytesIO,
    group_var: str = None,
    descriptive_stats: Dict[str, Any] = None,
    instrument_validation: Dict[str, Any] = None,
    time_series_results: Dict[str, Any] = None
) -> str:
    """
    Generates a structured Word report (.docx) and saves it with a readable filename.
    """
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    c_primary = RGBColor(30, 58, 138)
    c_text = RGBColor(31, 41, 55)
    
    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("LAPORAN HASIL ANALISIS DATA OTOMATIS")
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = c_primary
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_text = f"Penelitian: {research_title}\nFile Data: {filename}\nDihasilkan oleh Sistem Olahin"
    subtitle_run = subtitle_p.add_run(sub_text)
    subtitle_run.font.name = 'Calibri'
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.italic = True
    subtitle_run.font.color.rgb = RGBColor(107, 114, 128)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    def add_section_header(text: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = c_primary

    # 1. Summary & Preprocessing
    add_section_header("1. Ringkasan & Pembersihan Data")
    p = doc.add_paragraph()
    p.add_run(
        f"Data awal terdeteksi sebanyak {preprocessing_log['initial_rows']} baris. "
        f"Sistem preprocessing otomatis melakukan penyesuaian tipe data dan pembersihan data kosong/outlier pada kolom terpilih:\n"
    )
    p.add_run(f"- Variabel Dependen (Y): '{dependent_var}'\n").bold = True
    if independent_vars:
        p.add_run(f"- Variabel Independen (X): '{', '.join(independent_vars)}'\n").bold = True
    if group_var:
        p.add_run(f"- Variabel Grup: '{group_var}'\n").bold = True
        
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Light Shading Accent 1'
    table.rows[0].cells[0].text = 'Kriteria Pemrosesan'
    table.rows[0].cells[1].text = 'Keterangan / Hasil'
    table.rows[1].cells[0].text = 'Baris Data Awal'
    table.rows[1].cells[1].text = str(preprocessing_log['initial_rows'])
    table.rows[2].cells[0].text = 'Baris Data Bersih (Final)'
    table.rows[2].cells[1].text = str(preprocessing_log['final_rows'])
    
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    
    if preprocessing_log["actions_taken"]:
        doc.add_paragraph().add_run("Tindakan pembersihan otomatis:").bold = True
        for action in preprocessing_log["actions_taken"]:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            p.add_run(action)
            
    # 2. Questionnaire Instument Test (if applicable)
    if instrument_validation and instrument_validation.get("is_applicable"):
        add_section_header("2. Hasil Uji Validitas & Reliabilitas Instrumen")
        val = instrument_validation
        
        p = doc.add_paragraph()
        p.add_run(f"Hasil Uji Reliabilitas Cronbach's Alpha = **{val['reliability']['cronbach_alpha']:.4f}** (Status: **{'RELIABEL' if val['reliability']['is_reliable'] else 'TIDAK RELIABEL'}**).\n\n").bold = True
        p.add_run("Rincian Uji Validitas Item (Corrected Item-Total Correlation):")
        
        items_dict = val["validity"]["items"]
        table_val = doc.add_table(rows=len(items_dict) + 1, cols=4)
        table_val.style = 'Light Shading Accent 1'
        hdr = table_val.rows[0].cells
        hdr[0].text = 'Item'
        hdr[1].text = 'Koefisien Korelasi'
        hdr[2].text = 'p-value'
        hdr[3].text = 'Status'
        
        for idx, (item, res) in enumerate(items_dict.items(), start=1):
            row = table_val.rows[idx].cells
            row[0].text = item
            row[1].text = f"{res['correlation_coefficient']:.4f}"
            row[2].text = f"{res['p_value']:.4f}"
            row[3].text = "VALID" if res["is_valid"] else "TIDAK VALID"
            if not res["is_valid"]:
                row[3].paragraphs[0].runs[0].font.bold = True
                
    # 3. Descriptive Stats
    if descriptive_stats:
        add_section_header("3. Statistik Deskriptif Variabel")
        table_desc = doc.add_table(rows=len(descriptive_stats) + 1, cols=7)
        table_desc.style = 'Light Shading Accent 1'
        hdr = table_desc.rows[0].cells
        hdr[0].text = 'Variabel'
        hdr[1].text = 'Mean'
        hdr[2].text = 'Median'
        hdr[3].text = 'Modus'
        hdr[4].text = 'Std Dev'
        hdr[5].text = 'Min'
        hdr[6].text = 'Max'
        
        for idx, (col, stats_data) in enumerate(descriptive_stats.items(), start=1):
            row = table_desc.rows[idx].cells
            row[0].text = col
            row[1].text = f"{stats_data['mean']:.4f}"
            row[2].text = f"{stats_data['median']:.4f}"
            row[3].text = f"{stats_data['mode']:.4f}"
            row[4].text = f"{stats_data['std']:.4f}"
            row[5].text = f"{stats_data['min']:.4f}"
            row[6].text = f"{stats_data['max']:.4f}"

    # 4. Normality Check
    if not time_series_results:
        add_section_header("4. Hasil Uji Asumsi Normalitas")
        table_norm = doc.add_table(rows=4, cols=2)
        table_norm.style = 'Light Shading Accent 1'
        table_norm.rows[0].cells[0].text = 'Metode Uji Normalitas'
        table_norm.rows[0].cells[1].text = normality_results['test_name']
        table_norm.rows[1].cells[0].text = 'Nilai Statistik'
        table_norm.rows[1].cells[1].text = f"{normality_results['statistic']:.4f}"
        table_norm.rows[2].cells[0].text = 'p-value'
        table_norm.rows[2].cells[1].text = f"{normality_results['p_value']:.4f}"
        table_norm.rows[3].cells[0].text = 'Kesimpulan'
        table_norm.rows[3].cells[1].text = "Berdistribusi Normal" if normality_results["is_normal"] else "Tidak Berdistribusi Normal"
        table_norm.rows[3].cells[1].paragraphs[0].runs[0].font.bold = True
        
    # 5. Statistical Hypothesis / Modeling
    model_type = analysis_results.get("model_type", time_series_results.get("model_type", "Model Analisis") if time_series_results else "Uji Statistik")
    add_section_header(f"5. Hasil Analisis Hipotesis ({model_type})")
    
    # Specific tables based on model type
    if model_type in ["Simple Linear Regression", "Multiple Linear Regression"]:
        coefs = analysis_results["coefficients"]
        table_reg = doc.add_table(rows=len(coefs) + 1, cols=5)
        table_reg.style = 'Light Shading Accent 1'
        hdr = table_reg.rows[0].cells
        hdr[0].text = 'Variabel'
        hdr[1].text = 'Koefisien'
        hdr[2].text = 'Std Error'
        hdr[3].text = 't-Stat'
        hdr[4].text = 'p-value'
        
        for idx, (var, stats_data) in enumerate(coefs.items(), start=1):
            row = table_reg.rows[idx].cells
            row[0].text = "Konstanta" if var == 'const' else var
            row[1].text = f"{stats_data['coefficient']:.4f}"
            row[2].text = f"{stats_data['std_err']:.4f}"
            row[3].text = f"{stats_data['t_statistic']:.4f}"
            row[4].text = f"{stats_data['p_value']:.4f}"
            if stats_data['is_significant']:
                row[4].paragraphs[0].runs[0].font.bold = True
                
    elif "Correlation" in model_type:
        res_dict = analysis_results["results"]
        table_corr = doc.add_table(rows=len(res_dict) + 1, cols=5)
        table_corr.style = 'Light Shading Accent 1'
        hdr = table_corr.rows[0].cells
        hdr[0].text = 'Variabel Independen'
        hdr[1].text = 'Koefisien Korelasi'
        hdr[2].text = 'p-value'
        hdr[3].text = 'Kekuatan'
        hdr[4].text = 'Arah'
        
        for idx, (var, stats_data) in enumerate(res_dict.items(), start=1):
            row = table_corr.rows[idx].cells
            row[0].text = var
            row[1].text = f"{stats_data['correlation_coefficient']:.4f}"
            row[2].text = f"{stats_data['p_value']:.4f}"
            row[3].text = stats_data['strength']
            row[4].text = stats_data['direction']
            if stats_data['is_significant']:
                row[2].paragraphs[0].runs[0].font.bold = True
                
    elif model_type in ["Independent Samples T-Test", "Mann-Whitney U Test", "Paired Samples T-Test", "Wilcoxon Signed-Rank Test"]:
        groups = analysis_results["groups"]
        metrics = analysis_results["metrics"]
        g1 = groups["group1"]
        g2 = groups["group2"]
        
        table_groups = doc.add_table(rows=3, cols=4)
        table_groups.style = 'Light Shading Accent 1'
        hdr = table_groups.rows[0].cells
        hdr[0].text = 'Kelompok'
        hdr[1].text = 'N'
        hdr[2].text = 'Mean / Median'
        hdr[3].text = 'Keterangan'
        
        # Populate
        val1 = g1["mean"] if "T-Test" in model_type else g1["median"]
        val2 = g2["mean"] if "T-Test" in model_type else g2["median"]
        
        r1 = table_groups.rows[1].cells
        r1[0].text = str(g1["label"])
        r1[1].text = str(g1["count"])
        r1[2].text = f"{val1:.4f}"
        r1[3].text = "Kelompok 1"
        
        r2 = table_groups.rows[2].cells
        r2[0].text = str(g2["label"])
        r2[1].text = str(g2["count"])
        r2[2].text = f"{val2:.4f}"
        r2[3].text = "Kelompok 2"
        
        doc.add_paragraph().paragraph_format.space_after = Pt(6)
        
        table_stat = doc.add_table(rows=4, cols=2)
        table_stat.style = 'Light Shading Accent 1'
        table_stat.rows[0].cells[0].text = 'Metode Pengujian'
        table_stat.rows[0].cells[1].text = model_type
        table_stat.rows[1].cells[0].text = 'Statistik Uji'
        table_stat.rows[1].cells[1].text = f"{metrics['statistic']:.4f}"
        table_stat.rows[2].cells[0].text = 'p-value'
        table_stat.rows[2].cells[1].text = f"{metrics['p_value']:.4f}"
        table_stat.rows[3].cells[0].text = 'Effect Size'
        table_stat.rows[3].cells[1].text = f"{metrics['effect_size_name']}: {metrics['effect_size']:.4f}"
        if metrics['is_significant']:
            table_stat.rows[2].cells[1].paragraphs[0].runs[0].font.bold = True
            
    elif "Mediation" in model_type:
        paths = analysis_results["paths"]
        sobel = analysis_results["sobel_test"]
        
        table_paths = doc.add_table(rows=5, cols=4)
        table_paths.style = 'Light Shading Accent 1'
        hdr = table_paths.rows[0].cells
        hdr[0].text = 'Jalur'
        hdr[1].text = 'Koefisien'
        hdr[2].text = 'p-value'
        hdr[3].text = 'Status'
        
        row_names = [("path_a", "Jalur a (X \u2192 M)"), ("path_b", "Jalur b (M \u2192 Y)"), 
                     ("path_c_prime", "Jalur c' (Efek Langsung)"), ("total_effect_c", "Jalur c (Efek Total)")]
        for idx, (key, label) in enumerate(row_names, start=1):
            row = table_paths.rows[idx].cells
            row[0].text = label
            row[1].text = f"{paths[key]['coefficient']:.4f}"
            row[2].text = f"{paths[key]['p_value']:.4f}"
            row[3].text = "Signifikan" if paths[key].get('is_significant', False) else "Tidak Signifikan"

        doc.add_paragraph().paragraph_format.space_after = Pt(6)
        doc.add_paragraph(f"Hasil Sobel Test: Z = **{sobel['statistic_z']:.4f}**, p-value = **{sobel['p_value']:.4f}** (Efek Tidak Langsung = {sobel['indirect_effect']:.4f})").bold = True

    # 6. Interpretasi Naratif
    add_section_header("6. Interpretasi Hasil Analisis")
    for paragraph in narrative.split("\n\n"):
        if paragraph.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            parts = paragraph.split("**")
            is_bold = False
            for part in parts:
                run = p.add_run(part)
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
                run.font.color.rgb = c_text
                if is_bold:
                    run.bold = True
                is_bold = not is_bold
                
    # 7. Visualisasi Grafik
    add_section_header("7. Visualisasi Grafik")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    plot_buf.seek(0)
    p.add_run().add_picture(plot_buf, width=Inches(5.2))
    
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_p.add_run("Gambar: Visualisasi Grafik Analisis Statistik")
    caption_run.font.name = 'Calibri'
    caption_run.font.size = Pt(9.5)
    caption_run.font.italic = True
    caption_run.font.color.rgb = RGBColor(107, 114, 128)
    
    # Save Report
    safe_title = clean_filename(research_title)
    report_filename = f"{safe_title}_{file_id}_report.docx"
    report_path = os.path.join(EXPORT_DIR, report_filename)
    doc.save(report_path)
    return report_filename


# --- REPORTLAB PDF GENERATOR ---
def create_pdf_report(
    file_id: str,
    research_title: str,
    filename: str,
    preprocessing_log: Dict[str, Any],
    normality_results: Dict[str, Any],
    analysis_results: Dict[str, Any],
    independent_vars: List[str],
    dependent_var: str,
    narrative: str,
    plot_buf: io.BytesIO,
    group_var: str = None,
    descriptive_stats: Dict[str, Any] = None,
    instrument_validation: Dict[str, Any] = None,
    time_series_results: Dict[str, Any] = None
) -> str:
    """
    Generates a professional PDF report using reportlab and saves it to EXPORT_DIR.
    """
    safe_title = clean_filename(research_title)
    pdf_filename = f"{safe_title}_{file_id}_report.pdf"
    pdf_path = os.path.join(EXPORT_DIR, pdf_filename)
    
    doc = SimpleDocTemplate(
        pdf_path,
        pagesize=letter,
        rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54
    )
    
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=18,
        leading=22,
        textColor=colors.HexColor('#1E3A8A'),
        alignment=1, # Center
        spaceAfter=15
    )
    
    subtitle_style = ParagraphStyle(
        'DocSubtitle',
        parent=styles['Normal'],
        fontName='Helvetica-Oblique',
        fontSize=10,
        leading=14,
        textColor=colors.HexColor('#4B5563'),
        alignment=1, # Center
        spaceAfter=25
    )
    
    h1_style = ParagraphStyle(
        'SectionHeader',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=12,
        leading=16,
        textColor=colors.HexColor('#1E3A8A'),
        spaceBefore=15,
        spaceAfter=8,
        keepWithNext=True
    )
    
    body_style = ParagraphStyle(
        'ReportBody',
        parent=styles['BodyText'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13.5,
        textColor=colors.HexColor('#1F2937'),
        spaceAfter=6
    )
    
    bullet_style = ParagraphStyle(
        'ReportBullet',
        parent=body_style,
        leftIndent=15,
        firstLineIndent=-10,
        spaceAfter=3
    )
    
    story = []
    
    # Title Page/Header
    story.append(Paragraph("LAPORAN HASIL ANALISIS DATA STATISTIK", title_style))
    story.append(Paragraph(f"Penelitian: <b>{research_title}</b><br/>File Data: {filename}<br/>Dihasilkan oleh Sistem Olahin", subtitle_style))
    
    # Table styling parameters
    t_style = TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1E3A8A')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 9),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 5),
        ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#F3F4F6')),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#E5E7EB')),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.HexColor('#1F2937')),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 8.5),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
    ])

    # 1. Summary of ingestion & Preprocessing
    story.append(Paragraph("1. Ringkasan & Pembersihan Data", h1_style))
    story.append(Paragraph(f"Dataset awal memiliki total <b>{preprocessing_log['initial_rows']}</b> baris data. Setelah pembersihan tipe data, outlier, dan missing values, diperoleh <b>{preprocessing_log['final_rows']}</b> baris data bersih.", body_style))
    
    prep_data = [
        ["Parameter Pemrosesan", "Hasil / Angka"],
        ["Jumlah Baris Awal", str(preprocessing_log['initial_rows'])],
        ["Jumlah Baris Final", str(preprocessing_log['final_rows'])]
    ]
    t1 = Table(prep_data, colWidths=[250, 150])
    t1.setStyle(t_style)
    story.append(t1)
    story.append(Spacer(1, 10))
    
    if preprocessing_log["actions_taken"]:
        story.append(Paragraph("Tindakan pembersihan otomatis:", ParagraphStyle('BoldText', parent=body_style, fontName='Helvetica-Bold')))
        for action in preprocessing_log["actions_taken"]:
            story.append(Paragraph(f"&bull; {action}", bullet_style))
    story.append(Spacer(1, 10))
    
    # 2. Validity & Reliability (Questionnaire)
    if instrument_validation and instrument_validation.get("is_applicable"):
        story.append(Paragraph("2. Hasil Uji Validitas & Reliabilitas Kuesioner", h1_style))
        val = instrument_validation
        alpha_val = val["reliability"]["cronbach_alpha"]
        rel_lbl = "RELIABEL" if val["reliability"]["is_reliable"] else "TIDAK RELIABEL"
        
        story.append(Paragraph(f"Uji Reliabilitas Cronbach's Alpha = <b>{alpha_val:.4f}</b> (Status: <b>{rel_lbl}</b>, batas minimal = 0.6).", body_style))
        story.append(Spacer(1, 5))
        
        items_dict = val["validity"]["items"]
        val_data = [["Item", "Korelasi Item-Total", "p-value", "Status"]]
        for item, res in items_dict.items():
            val_data.append([
                item,
                f"{res['correlation_coefficient']:.4f}",
                f"{res['p_value']:.4f}",
                "VALID" if res["is_valid"] else "TIDAK VALID"
            ])
        t_val = Table(val_data, colWidths=[120, 100, 90, 90])
        t_val.setStyle(t_style)
        story.append(t_val)
        story.append(Spacer(1, 10))
        
    # 3. Descriptive Stats
    if descriptive_stats:
        story.append(Paragraph("3. Statistik Deskriptif Variabel", h1_style))
        desc_data = [["Variabel", "Mean", "Median", "Modus", "Std Dev", "Min", "Max"]]
        for col, stats_data in descriptive_stats.items():
            desc_data.append([
                col,
                f"{stats_data['mean']:.4f}",
                f"{stats_data['median']:.4f}",
                f"{stats_data['mode']:.4f}",
                f"{stats_data['std']:.4f}",
                f"{stats_data['min']:.4f}",
                f"{stats_data['max']:.4f}"
            ])
        t_desc = Table(desc_data, colWidths=[90, 50, 50, 50, 50, 50, 60])
        t_desc.setStyle(t_style)
        story.append(t_desc)
        story.append(Spacer(1, 10))
        
    # 4. Normality Check
    if not time_series_results:
        story.append(Paragraph("4. Hasil Uji Asumsi Normalitas", h1_style))
        norm_lbl = "Berdistribusi Normal" if normality_results["is_normal"] else "Tidak Berdistribusi Normal"
        norm_data = [
            ["Parameter", "Keterangan"],
            ["Metode Uji Normalitas", normality_results['test_name']],
            ["Statistik Uji", f"{normality_results['statistic']:.4f}"],
            ["p-value", f"{normality_results['p_value']:.4f}"],
            ["Kesimpulan Distribusi", norm_lbl]
        ]
        t_norm = Table(norm_data, colWidths=[200, 200])
        t_norm.setStyle(t_style)
        story.append(t_norm)
        story.append(Spacer(1, 10))
        
    # 5. Statistical Hypothesis
    model_type = analysis_results.get("model_type", time_series_results.get("model_type", "Uji Statistik") if time_series_results else "Uji Statistik")
    story.append(Paragraph(f"5. Uji Hipotesis & Pemodelan ({model_type})", h1_style))
    
    if model_type in ["Simple Linear Regression", "Multiple Linear Regression"]:
        coefs = analysis_results["coefficients"]
        reg_data = [["Variabel", "Koefisien (\u03b2)", "Std Error", "t-Statistik", "p-value"]]
        for var, stats_data in coefs.items():
            reg_data.append([
                "Konstanta" if var == 'const' else var,
                f"{stats_data['coefficient']:.4f}",
                f"{stats_data['std_err']:.4f}",
                f"{stats_data['t_statistic']:.4f}",
                f"{stats_data['p_value']:.4f}"
            ])
        t_reg = Table(reg_data, colWidths=[120, 70, 70, 70, 70])
        t_reg.setStyle(t_style)
        story.append(t_reg)
        
        # Summary
        metrics = analysis_results["metrics"]
        story.append(Spacer(1, 6))
        story.append(Paragraph(f"Model Fit: R-Square (R<sup>2</sup>) = <b>{metrics['r_squared']:.4f}</b>, F-statistic = <b>{metrics['f_statistic']:.4f}</b> (p-value = <b>{metrics['p_value']:.4f}</b>)", body_style))
        
    elif "Correlation" in model_type:
        res_dict = analysis_results["results"]
        corr_data = [["Variabel", "Koefisien Korelasi", "p-value", "Kekuatan", "Arah"]]
        for var, stats_data in res_dict.items():
            corr_data.append([
                var,
                f"{stats_data['correlation_coefficient']:.4f}",
                f"{stats_data['p_value']:.4f}",
                stats_data['strength'],
                stats_data['direction']
            ])
        t_corr = Table(corr_data, colWidths=[110, 80, 70, 70, 70])
        t_corr.setStyle(t_style)
        story.append(t_corr)
        
    elif model_type in ["Independent Samples T-Test", "Mann-Whitney U Test", "Paired Samples T-Test", "Wilcoxon Signed-Rank Test"]:
        groups = analysis_results["groups"]
        metrics = analysis_results["metrics"]
        g1 = groups["group1"]
        g2 = groups["group2"]
        val1 = g1["mean"] if "T-Test" in model_type else g1["median"]
        val2 = g2["mean"] if "T-Test" in model_type else g2["median"]
        metric_lbl = "Mean" if "T-Test" in model_type else "Median"
        
        g_data = [
            ["Kelompok", "N", f"Nilai Pusat ({metric_lbl})", "Deskripsi"],
            [g1["label"], str(g1["count"]), f"{val1:.4f}", "Kelompok 1"],
            [g2["label"], str(g2["count"]), f"{val2:.4f}", "Kelompok 2"]
        ]
        t_g = Table(g_data, colWidths=[120, 60, 100, 120])
        t_g.setStyle(t_style)
        story.append(t_g)
        story.append(Spacer(1, 6))
        
        story.append(Paragraph(f"Nilai Statistik Uji = <b>{metrics['statistic']:.4f}</b>, p-value = <b>{metrics['p_value']:.4f}</b> (Status: {'Signifikan' if metrics['is_significant'] else 'Tidak Signifikan'}).", body_style))
        story.append(Paragraph(f"Kekuatan efek ({metrics['effect_size_name']}) = <b>{metrics['effect_size']:.4f}</b>.", body_style))
        
    elif "Mediation" in model_type:
        paths = analysis_results["paths"]
        sobel = analysis_results["sobel_test"]
        
        path_data = [
            ["Jalur Pengaruh", "Koefisien arah", "p-value", "Signifikansi"],
            ["Jalur a (X \u2192 M)", f"{paths['path_a']['coefficient']:.4f}", f"{paths['path_a']['p_value']:.4f}", "Ya" if paths['path_a']['is_significant'] else "Tidak"],
            ["Jalur b (M \u2192 Y)", f"{paths['path_b']['coefficient']:.4f}", f"{paths['path_b']['p_value']:.4f}", "Ya" if paths['path_b']['is_significant'] else "Tidak"],
            ["Jalur c' (Efek Langsung)", f"{paths['path_c_prime']['coefficient']:.4f}", f"{paths['path_c_prime']['p_value']:.4f}", "Ya" if paths['path_c_prime']['is_significant'] else "Tidak"],
            ["Jalur c (Efek Total)", f"{paths['path_c']['coefficient']:.4f}", f"{paths['path_c']['p_value']:.4f}", "Ya" if paths['path_c']['is_significant'] else "Tidak"]
        ]
        t_paths = Table(path_data, colWidths=[130, 80, 80, 110])
        t_paths.setStyle(t_style)
        story.append(t_paths)
        story.append(Spacer(1, 6))
        
        story.append(Paragraph(f"Efek Mediasi (Sobel Test): Z = <b>{sobel['statistic_z']:.4f}</b>, p-value = <b>{sobel['p_value']:.4f}</b>. Nilai Efek Tidak Langsung = <b>{sobel['indirect_effect']:.4f}</b>.", body_style))
        
    story.append(Spacer(1, 10))
    story.append(PageBreak()) # Clean page break for narrative & visuals

    # 6. Interpretasi Naratif
    story.append(Paragraph("6. Interpretasi Hasil Analisis", h1_style))
    for paragraph in narrative.split("\n\n"):
        if paragraph.strip():
            # Format Markdown bold to reportlab html style
            formatted_p = paragraph.replace("**", "<b>", 1).replace("**", "</b>", 1)
            # Repeat to handle multiple bolds in a paragraph
            while "**" in formatted_p:
                formatted_p = formatted_p.replace("**", "<b>", 1).replace("**", "</b>", 1)
                
            story.append(Paragraph(formatted_p, body_style))
            story.append(Spacer(1, 4))
            
    story.append(Spacer(1, 10))
    
    # 7. Visualisasi Grafik
    story.append(Paragraph("7. Visualisasi Grafik", h1_style))
    # Write image to temp file in EXPORT_DIR for Reportlab to read
    temp_img_path = os.path.join(EXPORT_DIR, f"{file_id}_temp_plot.png")
    with open(temp_img_path, "wb") as f:
        f.write(plot_buf.getvalue())
        
    img = Image(temp_img_path, width=320, height=220)
    story.append(img)
    story.append(Spacer(1, 5))
    story.append(Paragraph("<i>Gambar: Hasil Plot Visualisasi Statistik Olahin</i>", ParagraphStyle('Caption', parent=body_style, alignment=1, fontSize=8.5, textColor=colors.HexColor('#6B7280'))))
    
    # Build Document
    try:
        doc.build(story)
    finally:
        # Clean up temp image
        if os.path.exists(temp_img_path):
            os.remove(temp_img_path)
            
    return pdf_filename
